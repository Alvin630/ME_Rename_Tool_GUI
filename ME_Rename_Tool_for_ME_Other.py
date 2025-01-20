import os
import tkinter as tk
from tkinter import ttk, filedialog, messagebox, simpledialog, Toplevel, Text
from PIL import Image, ImageTk
from tkinter.font import Font
import pprint
import platform
import subprocess
import shutil
from tkinterdnd2 import TkinterDnD, DND_FILES
from docx import Document
from docx.shared import Inches
import markdown

class CustomSimpledialog(simpledialog.Dialog):
    def __init__(self, parent, title, message, initialvalue=""):
        self.message = message
        self.initialvalue = initialvalue
        super().__init__(parent, title)

    def body(self, master):
        tk.Label(master, text=self.message).pack()
        self.entry = tk.Entry(master, width=50)
        self.entry.pack(padx=40, pady=10)
        self.entry.insert(0, self.initialvalue)
        return self.entry

    def apply(self):
        self.result = self.entry.get()

class FileExplorerApp:
    def __init__(self, root):
        self.root = root
        self.root.title("[ME] Other Test Data Rename Tool V-1.0.2")

        self.large_font = Font(family="Helvetica", size=12, weight="bold")

        # Frame for the first row of buttons
        self.frame1 = tk.Frame(self.root)
        self.frame1.pack(anchor="w",pady=5)

        # Frame for the second row of buttons
        self.frame2 = tk.Frame(self.root)
        self.frame2.pack(anchor="w",pady=5)

        self.create_folder = tk.Button(self.frame1, text="Create Folder",bg="#F9F900", command=self.trigger_folder_creation)
        self.create_folder.pack(side=tk.LEFT, padx=15)

        self.load_button = tk.Button(self.frame1, text="Load Directory", command=self.load_directory)
        self.load_button.pack(side=tk.LEFT, padx=15)

        self.expand_button = tk.Button(self.frame2, text="【+】", command=self.expand_all)
        self.expand_button.pack(side=tk.LEFT, padx=5)

        self.shrink_button = tk.Button(self.frame2, text="【-】", command=self.shrink_all)
        self.shrink_button.pack(side=tk.LEFT, padx=5)

        self.delete_button = tk.Button(self.frame2, text="Delete", command=self.delete_selected)
        self.delete_button.pack(side=tk.LEFT, padx=15)
        self.open_button = tk.Button(self.frame2, text="Open File", command=self.open_selected)
        self.open_button.pack(side=tk.LEFT, padx=15)
        
        self.rename_button = tk.Button(self.frame2, text="Rename Images in All Folders", command=self.rename_images)
        self.rename_button.pack(side=tk.LEFT, padx=15)

        # Add a new button for the Word processing functionality
        self.word_button = tk.Button(self.frame2, text="Insert Images into Word", command=self.insert_images_into_word)
        self.word_button.pack(side=tk.LEFT, padx=15)

        self.path_label = tk.Label(self.frame1, text="", font=("Helvetica", 14))
        self.path_label.pack(side=tk.LEFT, padx=5)

        self.folder_tree_frame = tk.Frame(root)
        self.folder_tree_frame.pack(side=tk.LEFT, fill=tk.Y)
        self.folder_tree_scroll = tk.Scrollbar(self.folder_tree_frame)
        self.folder_tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.folder_tree = ttk.Treeview(self.folder_tree_frame, yscrollcommand=self.folder_tree_scroll.set, height=20, style="Folder.Treeview")
        self.folder_tree.pack(side=tk.LEFT, fill=tk.Y)
        self.folder_tree_scroll.config(command=self.folder_tree.yview)

        self.image_tree_frame = tk.Frame(root)
        self.image_tree_frame.pack(side=tk.RIGHT, expand=True, fill=tk.BOTH)
        self.image_tree_scroll = tk.Scrollbar(self.image_tree_frame)
        self.image_tree_scroll.pack(side=tk.RIGHT, fill=tk.Y)
        self.image_tree = ttk.Treeview(self.image_tree_frame, yscrollcommand=self.image_tree_scroll.set, height=20, style="Image.Treeview")
        self.image_tree.pack(side=tk.LEFT, expand=True, fill=tk.BOTH)
        self.image_tree_scroll.config(command=self.image_tree.yview)

        self.folder_tree.bind("<ButtonRelease-1>", self.on_folder_click)
        self.image_tree.bind("<Double-1>", self.on_double_click)

        # Enable drag-and-drop in the folder tree
        self.root.drop_target_register(DND_FILES)
        self.root.dnd_bind('<<Drop>>', self.on_drop)

        self.dir_path = ""
        self.image_cache = {}

        self.folder_tree.column("#0", width=550)
        self.image_tree.column("#0", width=450)

        style = ttk.Style()
        style.configure('Treeview', rowheight=20)
        style.configure('Folder.Treeview', rowheight=50)
        style.configure('Image.Treeview', rowheight=300)

        self.folder_tree.tag_configure('folder', font=self.large_font)
        self.image_tree.tag_configure('image', font=self.large_font)
        self.image_tree.tag_configure('file', font=self.large_font)
        # Assuming `self.folder_tree` is your Treeview
        self.folder_tree.tag_configure('folder', foreground='blue',font=('Helvetica', 12, 'bold'))  # Folder items will be yellow
        self.folder_tree.tag_configure('file', foreground='black',font=('Helvetica', 10, 'bold'))   # File items will be black

#要討論六面是否要按照順序，如按照順序可限定naming pattern 順序
        self.default_naming_patterns = {
            # IP2X, IP3X, IP4X
            "IP 2X Basic Function Test": {
                "": [
                    "IP 2X Function Check Before Test", 
                    "IP 2X Function Check After Test"
                ]
            },
            "IP 2X Test": {
                "": ["IP 2X Test"]
            },
            "IP 3X Basic Function Test": {
                "": [
                    "IP 3X Function Check Before Test", 
                    "IP 3X Function Check After Test"
                ]
            },
            "IP 3X Test": {
                "": ["IP 3X Test"]
            },
            "IP 4X Basic Function Test": {
                "": [
                    "IP 4X Function Check Before Test", 
                    "IP 4X Function Check After Test"
                ]
            },
            "IP 4X Test": {
                "": ["IP 4X Test"]
            },

            # IP5X, IP6X, IPX5, IPX6, IPX7, IPX8
            "IP 5X Basic Function Test": {
                "": [
                    "IP 5X Function Check Before Test", 
                    "IP 5X Function Check After Test"
                ]
            },
            "IP 5X Test": {
                "": ["IP 5X Test"]
            },
            "IP 66 Basic Function Test": {
                "": [
                    "IP 66 Function Check Before Test", 
                    "IP 66 Function Check After Test"
                ]
            },
            "IP 6X Test": {
                "": ["IP 6X Test"]
            },
            "IP X6 Test": {
                "": ["IP X6 Test"]
            },
            "IP66老化測試": {
                "": ["IP 66 Setup", "IP66 Finish"]
            },
            "IP 67 Basic Function Test": {
                "": [
                    "IP 67 Function Check Before Test", 
                    "IP 67 Function Check After Test"
                ]
            },
            "IP X7 Test": {
                "": ["IP X7 Test"]
            },
            "IP67老化測試": {
                "": ["IP 67 Setup", "IP67 Finish"]
            },
            "IP 68 Basic Function Test": {
                "": [
                    "IP 68 Function Check Before Test", 
                    "IP 68 Function Check After Test"
                ]
            },
            "IP X8 Test": {
                "": ["IP X8 Test"]
            },
            "IP68老化測試": {
                "": ["IP 68 Setup", "IP68 Finish"]
            },

            # Salt Mist
            "Salt Mist Basic Function Test": {
                "": [
                    "Function Check Before Test", 
                    "Salt Mist Function Check After Test"
                ]
            },
            "Salt Mist Test": {
                "": [
                    "Salt Mist Function Check Before Test", 
                    "Test-Start", 
                    "Test-Finish", 
                    "Salt Mist Function Check After Test"
                ]
            },

            # Package (PKG)
            "Step3_Package Basic Test": {
                "": [
                    "PKG Function Check Before Test", 
                    "PKG Function Check After Test"
                ]
            },
            "Step4_Package Vib Test": {
                "": ["Z-Start", "Z-Finish"]
            },
            "Step6_Package Drop Test": {
                "": [
                    "weight_height", 
                    "Corner", 
                    "Edge1", 
                    "Edge2", 
                    "Edge3", 
                    "Face1", 
                    "Face2", 
                    "Face3", 
                    "Face4", 
                    "Face5", 
                    "Face6", 
                    "Unboxing"
                ]
            },
            "Step9_Function Check After ISTA Test": {
                "": ["Function check"]
            },

            # Storage
            "Storage Test Basic": {
                "": ["Function check Before Test", "Function check After Test"]
            },
            "HTHHS Test": {
                "": ["Setup", "Finish", "Function check After Test"]
            },
            "LTS Test": {
                "": ["Setup", "Finish", "Function check After Test"]
            }
        }
        if os.path.exists(default_md_file):
            self.show_markdown_in_new_window(default_md_file)
        else:
            messagebox.showerror("Error", f"Default Markdown file not found:\n{default_md_file}")

    #定義markdown window auto open
    def show_markdown_in_new_window(self, file_path):
        # Read and convert the Markdown file to HTML
        try:
            with open(file_path, 'r', encoding='utf-8') as file:
                md_content = file.read()
            html_content = markdown.markdown(md_content)
        except Exception as e:
            messagebox.showerror("Error", f"Failed to read the file: {e}")
            return

        # Create a new Toplevel window
        new_window = Toplevel(self.root)
        new_window.title(f"Markdown Viewer - {os.path.basename(file_path)}")

        # Add a scrollable Text widget to display the content
        text_widget = tk.Text(new_window, wrap=tk.WORD)
        text_widget.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)

        # Insert the HTML (rendered Markdown) into the Text widget
        text_widget.insert(tk.END, html_content)
        text_widget.configure(state="disabled")  # Make the content read-only

#創建資料夾結構函數
#創建資料夾結構       
# Function to create the folder structure
    def create_folder_structure(self, test_case_name, selected_executions, sample_sixpiece, base_path):
        clean_test_case_name = test_case_name.strip()

        # Create the outer folder for the test case
        dir_path = os.path.join(base_path, clean_test_case_name)
        os.makedirs(dir_path, exist_ok=True)

        # Iterate over each selected test execution
        for test_execution in selected_executions:
            test_execution_names = self.get_test_execution_names(test_execution)

            for exec_name in test_execution_names:
                exec_dir = os.path.join(dir_path, exec_name)
                os.makedirs(exec_dir, exist_ok=True)

                # elif exec_name in ["Package Vib Test"]:
                #     dut_test_dir = os.path.join(exec_dir, "外箱六面_Vib前","彩盒&Test DUT六面_Vib前")
                #     os.makedirs(dut_test_dir, exist_ok=True)
                #==================================================================
                # Check if the folder is named "Test DUT 六面" and create subfolders
                if "DUT六面" in exec_name:
                    if sample_sixpiece and sample_sixpiece > 1:
                        for i in range(1, sample_sixpiece + 1):
                            burn_in_sample_dir = os.path.join(exec_dir, str(i))
                            os.makedirs(burn_in_sample_dir, exist_ok=True)
                #==================================================================
                            
        # Notify the user of success and open the folder
        messagebox.showinfo("Success", f"Folder structure created at: {dir_path}")
        # Set the dir_path to the newly created directory
        self.dir_path = dir_path
        # Update the tree view to reflect the newly created folder structure
        self.refresh_views()
        # Select the created folder in the tree view
        self.reselect_folder(dir_path)
        # Open the folder in the file explorer (optional, if you want to automatically open it)
        self.open_folder(dir_path)

        # Update the path label
        self.path_label.config(text="路徑：" + dir_path)


    #定義test execution name 函數
    def get_test_execution_names(self, test_execution):
        """Helper function to map test_execution to actual folder names."""
        if test_execution == "[IP dirt] IP 2X":
            return [
                "IP 2X Basic Function Test",
                "IP 2X Test DUT六面",
                "IP 2X Test"
            ]
        elif test_execution == "[IP dirt] IP 3X":
            return [
                "IP 3X Basic Function Test",
                "IP 3X Test DUT六面",
                "IP 3X Test"
            ]
        elif test_execution == "[IP dirt] IP 4X":
            return [
                "IP 4X Basic Function Test",
                "IP 4X Test DUT六面",
                "IP 4X Test"
            ]
        elif test_execution == "[IP dirt] IP 5X":
            return [
                "IP 5X Basic Function Test",
                "IP 5X Test DUT六面",
                "IP 5X Test"
            ]
        elif test_execution == "[IP dirt] IP 6X":
            return [
                "IP 6X Basic Function Test",
                "IP 6X Test DUT六面",
                "IP 6X Test"
            ]
        elif test_execution == "[IP Water] IP X3":
            return [
                "IP X3 Basic Function Test",
                "IP X3 Test DUT六面",
                "IP X3 Test"
            ]
        elif test_execution == "[IP Water] IP X4":
            return [
                "IP X4 Basic Function Test",
                "IP X4 Test DUT六面",
                "IP X4 Test"
            ]
        elif test_execution == "[IP Water] IP X5":
            return [
                "IP X5 Basic Function Test",
                "IP X5 Test DUT六面",
                "IP X5 Test"
            ]
        elif test_execution == "[IP Water] IP X6":
            return [
                "IP X6 Basic Function Test",
                "IP X6 Test DUT六面",
                "IP X6 Test"
            ]
        elif test_execution == "[IP Water] IP X7":
            return [
                "IP X7 Basic Function Test",
                "IP X7 Test DUT六面",
                "IP X7 Test"
            ]
        elif test_execution == "[IP Water] IP X8":
            return [
                "IP X8 Basic Function Test",
                "IP X8 Test DUT六面",
                "IP X8 Test"
            ]
        elif test_execution == "[IP Test] IP 66":
            return [
                "IP 66 Basic Function Test",
                "IP 6X Test DUT六面",
                "IP X6 Test DUT六面",
                "IP 6X Test",
                "IP X6 Test",
                "IP66老化測試"
            ]
        elif test_execution == "[IP Test] IP 67":
            return [
                "IP 67 Basic Function Test",
                "IP 6X Test DUT六面",
                "IP X7 Test DUT六面",
                "IP 6X Test",
                "IP X7 Test",
                "IP67老化測試"
            ]
        elif test_execution == "[IP Test] IP 68":
            return [
                "IP 68 Basic Function Test",
                "IP 6X Test DUT六面",
                "IP X8 Test DUT六面",
                "IP 6X Test",
                "IP X8 Test",
                "IP68老化測試"
            ]
        elif test_execution == "[Salt Mist]":
            return [
                "Salt Mist Test 測試前DUT六面",
                "Salt Mist Test",
                "Salt Mist Test 測試後DUT六面"
            ]
        
        elif test_execution == "[ISTA] Package Vib&Drop":
            return [
                "Step1_外箱六面_Vib前",
                "Step2_彩盒&DUT六面_Vib前",
                "Step3_Package Basic Test",
                "Step4_Package Vib Test",
                "Step5_外箱六面_Vib後Drop前",
                "Step6_Package Drop Test",
                "Step7_外箱六面_Drop後",
                "Step8_彩盒&DUT六面_Drop",
                "Step9_Function Check After ISTA Test"
            ]
        #需新增Storage Basic function test folder
        elif test_execution == "[Storage] HTHHS&LTS":
            return [
                "Storage Test Basic",
                "Storage Test Basic DUT六面"
                "HTHHS Test",
                "After HTHHS Test DUT六面",
                "LTS Test",
                "After LTS Test DUT六面"
            ]
        else:
            return [test_execution]

#創建資料夾函數
    def open_folder(self, path):
        # Detect the operating system and open the folder accordingly
        if platform.system() == "Windows":
            os.startfile(path)
        elif platform.system() == "Darwin":  # macOS
            subprocess.Popen(["open", path])
        elif platform.system() == "Linux":
            subprocess.Popen(["xdg-open", path])
        else:
            messagebox.showerror("Error", "Unsupported operating system.")

#定義使用者輸入並創建資料夾函數
    def get_user_input_and_create_folders(self):
        # Helper functions to get valid inputs
        def ask_valid_string(prompt_title, prompt_text):
            while True:
                user_input = simpledialog.askstring(prompt_title, prompt_text, parent=self.root)
                if user_input is None:
                    messagebox.showwarning("Input Cancelled", "使用者取消!  ")
                    return None
                elif user_input.strip() == "":
                    messagebox.showwarning("Invalid Input", "此為必填，請輸入有效字元!")
                else:
                    return user_input

        def ask_valid_integer(prompt_title, prompt_text):
            while True:
                user_input = simpledialog.askinteger(prompt_title, prompt_text, parent=self.root)
                if user_input is None:
                    messagebox.showwarning("Input Cancelled", "使用者取消!  ")
                    return None
                else:
                    return user_input
                
        def select_test_executions():
            selected_executions = []
            options = [ 
                        "[IP dirt] IP 2X", "[IP dirt] IP 3X","[IP dirt] IP 4X",
                        "[IP dirt] IP 5X", "[IP dirt] IP 6X","[IP water] IP X3", "[IP water] IP X4",
                        "[IP water] IP X5","[IP water] IP X6", "[IP water] IP X7","[IP water] IP X8",
                        "[IP Test] IP 66", "[IP Test] IP 67", "[IP Test] IP 68", "[Salt Mist]", 
                        "[ISTA] Package Vib&Drop", 
                        "[Storage] HTHHS&LTS"
                        ]
            #Function to handle the window closure behavior (close button X)
            # Create a new window for the multi-select dropdown menu
            execution_window = tk.Toplevel(self.root)
            execution_window.title("Test Executions")
            execution_window.attributes("-topmost", True)

            tk.Label(execution_window, text="選擇欲測試之ME測項:").pack(pady=10)

            # Multi-select dropdown menu for test execution
            test_executions = tk.Listbox(execution_window, selectmode=tk.MULTIPLE, height=20, width=40)
            for item in options:
                test_executions.insert(tk.END, item)
            test_executions.pack(pady=10)

            def on_close():
                """Handle when the user closes the window via the 'X' button."""
                messagebox.showwarning("Input Cancelled", "使用者取消!")
                execution_window.destroy()  # Close the window
                return None  # Ensure it returns None when the window is closed

            # Adding the close button functionality (X)
            execution_window.protocol("WM_DELETE_WINDOW", on_close)

            def confirm_executions():
                selected_indices = test_executions.curselection()
                for i in selected_indices:
                    selected_executions.append(options[i])  # Get the actual string values
                if not selected_executions:
                    messagebox.showwarning("Invalid Input", "此為必選項目,請選擇測項!")
                    
                else:
                    execution_window.destroy()

            tk.Button(execution_window, text="Enter", command=confirm_executions).pack(pady=10)

            # Wait until the window is closed
            execution_window.wait_window()

            return selected_executions if selected_executions else None

        # Get test case name
        test_case_name = ask_valid_string("Test Case Name", "請輸入測試專案號碼+測試樣品名(Jira)!     ").strip()
        if test_case_name is None:
            return

        # # Get test sample and chamber model names (allow skipping)
        # test_sample_chamber_names = test_sample_chamber()

        # # If the user skipped the sample input, continue with an empty list
        # if test_sample_chamber_names is None:
        #     return

        # Get test executions
        selected_executions = select_test_executions()
        if selected_executions is None:
            return

        # Get burn-in samples
        samples_sixpiece = ask_valid_integer("DUT六面樣品數量", "請輸入DUT六面樣品紀錄總數量!")
        if samples_sixpiece is None:
            return

        # Ask the user to choose a directory where the structure will be created
        base_path = filedialog.askdirectory(title="Select Base Directory")
        if not base_path:
            messagebox.showerror("Error", "請選擇創建資料夾位置!       ")
            return

        # Create the folder structure
        self.create_folder_structure(test_case_name, selected_executions, samples_sixpiece, base_path)


        
    #定義TK GUI 創建資料夾函數按鈕功能函數
    def trigger_folder_creation(self):
        # Trigger the folder creation process
        self.get_user_input_and_create_folders()


    def load_directory(self):
        self.dir_path = filedialog.askdirectory()
        if self.dir_path:
            self.path_label.config(text="路徑：" + self.dir_path)
            self.folder_tree.delete(*self.folder_tree.get_children())
            self.image_tree.delete(*self.image_tree.get_children())
            self.load_folders(self.dir_path)

    #新增左邊資料夾區之內部檔案表現
    def load_folders(self, path, parent=""):
        # First, store image files and other files separately
        image_files = []
        other_files = []
        folders = []
        
        for item in os.listdir(path):
            abs_path = os.path.join(path, item)

            if os.path.isdir(abs_path):
                # Store the folder for later insertion
                folders.append(item)
            else:
                # If the item is a file, categorize it
                file_extension = os.path.splitext(item)[1].lower()
                if file_extension in ['.png', '.jpg', '.jpeg']:
                    image_files.append(item)
                else:
                    other_files.append(item)

        # Now insert image files first
        for image in image_files:
            self.folder_tree.insert(parent, 'end', text=image, tags=('file',))
        
        # Then insert the other files
        for file in other_files:
            self.folder_tree.insert(parent, 'end', text=file, tags=('file',))
        
        # Finally, insert folders and load their contents recursively
        for folder in folders:
            node = self.folder_tree.insert(parent, 'end', text=folder, open=False, tags=('folder',))
            self.load_folders(os.path.join(path, folder), node)

    def on_folder_click(self, event):
        # Get the current selection from the folder tree
        selection = self.folder_tree.selection()

        # Ensure that the selection is not empty to avoid IndexError
        if selection:
            selected_item = selection[0]
            folder_path = self.get_item_path(selected_item)

            # Check if the selected item is a folder
            if os.path.isdir(folder_path):
                self.load_files_and_images(folder_path)
            else:
                # Deselect the file item to prevent any action
                self.folder_tree.selection_remove(selected_item)
        else:
            print("No folder selected.")

    def load_files_and_images(self, path):
        self.image_tree.delete(*self.image_tree.get_children())
        for item in os.listdir(path):
            abs_path = os.path.join(path, item)
            if os.path.isfile(abs_path):
                file_extension = os.path.splitext(item)[1].lower()
                if file_extension in ['.png', '.jpg', '.jpeg']:
                    if abs_path not in self.image_cache:
                        img = Image.open(abs_path)
                        img.thumbnail((370, 370))
                        self.image_cache[abs_path] = ImageTk.PhotoImage(img)
                    image = self.image_cache[abs_path]
                    self.image_tree.insert('', 'end', text=item, image=image, tags=('image',))
                else:
                    self.image_tree.insert('', 'end', text=item, tags=('file',))
    #用在建構Folder tree
    def get_item_path(self, item_id):
        path_components = []
        while item_id:
            item_text = self.folder_tree.item(item_id, "text")
            path_components.append(item_text)
            item_id = self.folder_tree.parent(item_id)
        path_components.reverse()
        return os.path.join(self.dir_path, *path_components)
    
    #更新樹狀圖介面(待修正錯誤問題)
    def refresh_views(self):
        # Reselect the previously selected folder with a slight delay
        def refresh_with_delay():
            # Remember the currently selected folder
            selected_folder_id = self.folder_tree.selection()
            if selected_folder_id:
                selected_folder_path = self.get_item_path(selected_folder_id[0])
            else:
                selected_folder_path = self.dir_path

            # Clear and reload the folder and image trees
            self.folder_tree.delete(*self.folder_tree.get_children())
            self.image_tree.delete(*self.image_tree.get_children())
            self.load_folders(self.dir_path)

            # Expand all the folders first
            self.expand_all()

            # Reselect the previously selected folder
            self.reselect_folder(selected_folder_path)

            # Reload the images in the reselected folder
            self.load_files_and_images(selected_folder_path)

        # Use self.after to introduce a small delay (e.g., 100 ms)
        self.root.after(100, refresh_with_delay)

    def reselect_folder(self, path):
        # Recursively search for the folder by path and select it
        def recursive_select(node, target_path):
            item_path = self.get_item_path(node)
            if item_path == target_path:
                self.folder_tree.selection_set(node)
                self.folder_tree.see(node)
                return True
            for child in self.folder_tree.get_children(node):
                if recursive_select(child, target_path):
                    return True
            return False

        for child in self.folder_tree.get_children(''):
            if recursive_select(child, path):
                break
    #雙擊重新命名功能
    def on_double_click(self, event):
        item_id = self.image_tree.selection()[0]
        item_text = self.image_tree.item(item_id, "text")
        item_path = self.get_image_item_path(item_id)
        if not item_path:
            return  # Stop if the path could not be retrieved

        original_extension = os.path.splitext(item_text)[1]

        new_name = CustomSimpledialog(self.root, title="Rename", message="請輸入新名稱：", initialvalue=item_text).result
        if new_name:
            new_extension = os.path.splitext(new_name)[1]
            if not new_extension:
                new_name += original_extension

            new_path = os.path.join(os.path.dirname(item_path), new_name)

            # Check if the new file name already exists
            if os.path.exists(new_path):
                messagebox.showwarning(
                    "檔案已存在",
                    f"檔案 '{new_name}' 已存在於資料夾. 重命名已取消."
                )
                return  # Exit to prevent renaming

            try:
                os.rename(item_path, new_path)
                self.refresh_views()
            except Exception as e:
                messagebox.showerror("Error", f"Failed to rename file: {e}")

    #用在建構image tree & file tree
    def get_image_item_path(self, item_id):
        # Check if any folder is selected in the folder_tree
        folder_selection = self.folder_tree.selection()
        if not folder_selection:
            # No folder is selected, handle it gracefully (you can show a message if needed)
            messagebox.showerror("Error", "No folder selected!(請選擇資料夾路徑!)")
            return None
        # If folder is selected, construct the path
        selected_folder = self.get_item_path(self.folder_tree.selection()[0])
        image_name = self.image_tree.item(item_id, "text")
        return os.path.join(selected_folder, image_name)
    
    #用於拖檔案(複製)至GUI介面中    
    def on_drop(self, event):
        # Get the destination folder from the treeview selection
        dest_item = self.folder_tree.selection()
        dest_path = self.get_item_path(dest_item[0]) if dest_item else self.dir_path

        if not os.path.isdir(dest_path):
            messagebox.showerror("Error", "請選擇有效的目標資料夾！")
            return

        dropped_files = self.root.tk.splitlist(event.data)
        for file in dropped_files:
            try:
                if os.path.isfile(file):
                    dest_file = os.path.join(dest_path, os.path.basename(file))
                    if os.path.exists(dest_file):
                        if not messagebox.askyesno("File Exists", f"檔案 {os.path.basename(file)} 已存在。是否要覆蓋？"):
                            continue
                    shutil.copy2(file, dest_path)
                elif os.path.isdir(file):
                    dest_dir = os.path.join(dest_path, os.path.basename(file))
                    if os.path.exists(dest_dir):
                        if not messagebox.askyesno("Folder Exists", f"資料夾 {os.path.basename(file)} 已存在。是否要覆蓋？"):
                            continue
                        shutil.rmtree(dest_dir)
                    shutil.copytree(file, dest_dir)
            except PermissionError:
                messagebox.showerror("Error", f"無權限存取檔案: {file}")
            except Exception as e:
                messagebox.showerror("Error", f"複製失敗: {str(e)}")
        
        self.load_files_and_images(dest_path)
        self.refresh_views()

    # 刪除資料夾或檔案功能函數
    def delete_selected(self):
        try:
            if self.image_tree.selection():  # Check if an image or file is selected
                selected_item = self.image_tree.selection()[0]
                file_path = self.get_image_item_path(selected_item)
                if messagebox.askyesno("Delete", "Are you sure you want to delete this file?"):
                    os.remove(file_path)
                    self.image_tree.delete(selected_item)  # Remove the item from the tree view
                    self.refresh_views()
            elif self.folder_tree.selection():  # Check if a folder is selected
                selected_item = self.folder_tree.selection()[0]
                folder_path = self.get_item_path(selected_item)
                if messagebox.askyesno("Delete", "Are you sure you want to delete this folder and its contents?"):
                    shutil.rmtree(folder_path)  # Delete the folder and its contents
                    self.folder_tree.delete(selected_item)  # Remove the item from the tree view
        except Exception as e:
            messagebox.showerror("Error", f"Failed to delete: {e}")

    #展開與摺疊
    def expand_all(self):
        for item in self.folder_tree.get_children():
            self.folder_tree.item(item, open=True)
            self.expand_all_children(item)

    def expand_all_children(self, item):
        for child in self.folder_tree.get_children(item):
            self.folder_tree.item(child, open=True)
            self.expand_all_children(child)

    def shrink_all(self):
        for item in self.folder_tree.get_children():
            self.folder_tree.item(item, open=False)
            self.shrink_all_children(item)

    def shrink_all_children(self, item):
        for child in self.folder_tree.get_children(item):
            self.folder_tree.item(child, open=False)
            self.shrink_all_children(child)

# 開啟檔案功能 
    def open_selected(self):
        selected_item = self.image_tree.selection()
        if selected_item:
            item_path = self.get_image_item_path(selected_item[0])
            if os.path.isfile(item_path):
                try:
                    os.startfile(item_path)
                except OSError as e:
                    messagebox.showerror("錯誤", f"無法開啟圖片 {item_path}: {e}")
            elif os.path.isdir(item_path):
                try:
                    os.startfile(item_path)
                except OSError as e:
                    messagebox.showerror("錯誤", f"無法開啟資料夾 {item_path}: {e}")

    #定義已存在之名稱與不符合naming rule 之名稱
    def validate_and_rename_images(self, folder_path,test_type):
        """
        Validate and rename images in the folder based on the naming patterns.
        Logic:
        - If all images match the naming patterns, prompt: "覆蓋已命名好的檔案嗎？".
        - If no images match, prompt: "所有影像不符合命名規則，是否重命名？".
        - If some images match and some don't:
            - Show existing names and unmatched names.
            - Prompt: "是否重命名此資料夾的影像？順序可能會亂。".
        """
        # Retrieve the naming patterns for the specific test type
        naming_patterns = self.default_naming_patterns.get(test_type, {}).get("", [])

        files = sorted(os.listdir(folder_path))
        images = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg'))]

        matched_files = set()
        mismatched_files = []
        existing_names = []

        # Classify images based on naming patterns
        for file_name in images:
            base_name, _ = os.path.splitext(file_name)

            if base_name in naming_patterns:
                matched_files.add(base_name)
                existing_names.append(file_name)
            else:
                mismatched_files.append(file_name)

        missing_patterns = [p for p in naming_patterns if p not in matched_files]

        if not mismatched_files and not missing_patterns:
            # Case 1: All images match naming patterns
            return self.prompt_for_overwrite(folder_path, existing_names)

        elif not matched_files:
            # Case 2: None of the images match naming patterns #已修正使用者提醒未命名
            return self.prompt_for_rename_all(folder_path, mismatched_files, missing_patterns, naming_patterns, images)

        else:
            # Case 3: Partial match
            return self.prompt_for_partial_rename(folder_path, existing_names, mismatched_files, missing_patterns, naming_patterns, images)
    
    #==============將命名路徑特別標記提醒的函數(待修改) =======================================
    # def show_highlight_popup(self, message):
    #     """
    #     Display a popup window with highlighted specific words in the message.
    #     """
    #     if not isinstance(message, str):  # Ensure the message is a string
    #         raise ValueError("Expected 'message' to be a string")

    #     # Create a new Toplevel window
    #     popup = tk.Toplevel(self.root)
    #     popup.title("Highlight Popup")
    #     popup.geometry("400x300")

    #     # Add a scrollable Text widget
    #     text_widget = tk.Text(popup, wrap=tk.WORD)
    #     text_widget.pack(expand=True, fill=tk.BOTH, padx=10, pady=10)

    #     # Insert the message line by line, with optional highlighting
    #     for line in message.split("\n"):
    #         if "已符合命名規則的檔案：" in line:
    #             text_widget.insert(tk.END, line + "\n", "highlight")
    #         else:
    #             text_widget.insert(tk.END, line + "\n")

    #     # Configure the Text widget to highlight specific text
    #     text_widget.tag_configure("highlight", foreground="red", font=("Arial", 12, "bold"))
    #     text_widget.configure(state="disabled")  # Make the text read-only
    
    # def prompt_for_overwrite(self, folder_path, existing_names):
    #     """
    #     Prompt the user to overwrite all existing names.
    #     """
    #     message = (
    #         f"資料夾：'{folder_path}' 的影像已符合命名規則。\n\n"
    #         f"已符合命名規則的檔案：\n{', '.join(existing_names)}\n\n"
    #         f"你要覆蓋已命名好的檔案嗎？"
    #     )

    #     # Display the highlighted popup
    #     self.show_highlight_popup(message)

    #     # Show a Yes/No messagebox after the popup
    #     response = tk.messagebox.askyesno("覆蓋確認", "你要覆蓋已命名好的檔案嗎？")
    #     return response
    #=======================================================================================

#嘗試修正對正確的資料夾檔案判定為錯誤 (問題為因原本的名稱就跟naming rule不同，系統會自行判斷不符合命名規則)
    def prompt_for_rename_all(self, folder_path,  mismatched_files, missing_patterns,naming_patterns, images):
        """
        Prompt the user to rename all files if none match.
        """
        response = messagebox.askyesno(
            "重命名建議",
            f"資料夾：'{folder_path}' 的所有影像不符合命名規則。\n\n" 
            f"不符合命名規則的檔案：\n{', '.join(mismatched_files)}\n\n"
            f"尚未滿足的命名規則：\n{', '.join(missing_patterns)}\n\n"
            f"是否要重命名？"
        )
        if response:
            return self.perform_rename(folder_path, naming_patterns, images)
        return 0
#重新檢查命盟規則是否符合naming pattern狀況
    def prompt_for_partial_rename(self, folder_path, existing_names, mismatched_files, missing_patterns, naming_patterns, images):
        """
        Prompt the user to handle partial matches.
        """
        warning_message = (
            f"在資料夾 '{folder_path}' 中發現以下情況：\n\n"
            f"已符合命名規則的檔案：\n{', '.join(existing_names)}\n\n"
            f"不符合命名規則的檔案：\n{', '.join(mismatched_files)}\n\n"
            f"尚未滿足的命名規則：\n{', '.join(missing_patterns)}\n\n"
            f"請檢查影像順序，否則順序可能會亂，是否要重新命名？"
        )
        response = messagebox.askyesno("重命名確認", warning_message)
        if response:
            return self.perform_rename(folder_path, naming_patterns, images)
        return 0

    def perform_rename(self, folder_path, naming_patterns, images):
        """
        Rename images based on naming patterns, handling duplicates.
        """
        used_names = set()
        renamed_count = 0

        for image_name in images:
            src = os.path.join(folder_path, image_name)
            base_name, file_extension = os.path.splitext(image_name)

            # Skip renaming if already matches the naming pattern
            if base_name in naming_patterns and base_name not in used_names:
                used_names.add(base_name)
                continue

            # Find an unused valid name for renaming
            new_base_name = next((name for name in naming_patterns if name not in used_names), None)
            if new_base_name is None:
                continue

            new_name = new_base_name + file_extension
            dst = os.path.join(folder_path, new_name)

            try:
                os.rename(src, dst)
                used_names.add(new_base_name)
                renamed_count += 1
            except Exception as e:
                messagebox.showerror("Error", f"重命名失敗: {e}")

        return renamed_count

    def check_and_rename_outer_layer(self):
        total_renamed = 0

        # Traverse directories in the main path
        outer_layer_folders = [
            d for d in os.listdir(self.dir_path) if os.path.isdir(os.path.join(self.dir_path, d))
        ]

        for folder in outer_layer_folders:
            folder_path = os.path.join(self.dir_path, folder)

            if folder in self.default_naming_patterns:
                # Validate the folder contents against naming patterns
                if not self.validate_and_rename_images(folder_path, folder):
                    continue  # Skip renaming if validation fails and user chooses not to pass

                # Proceed with renaming after validation
                total_renamed += self.process_test_type_folder_with_check(self.dir_path, folder)

            else:
                for test_type, subfolders in self.default_naming_patterns.items():
                    test_type_path = os.path.join(folder_path, test_type)

                    if os.path.isdir(test_type_path):
                        # Validate the folder contents against naming patterns
                        if not self.validate_and_rename_images(test_type_path, test_type):
                            continue  # Skip renaming if validation fails and user chooses not to pass

                        # Proceed with renaming after validation
                        total_renamed += self.process_test_type_folder_with_check(folder_path, test_type)

        return total_renamed

    def process_test_type_folder_with_check(self, base_path, test_type):
        """
        Process images inside a test_type folder by renaming them according to the naming patterns,
        with checks for existing names.
        """
        total_renamed = 0
        subfolders = self.default_naming_patterns.get(test_type, {})

        for subfolder, naming_patterns in subfolders.items():
            folder_path = os.path.join(base_path, test_type, subfolder)
            if os.path.isdir(folder_path):
                total_renamed += self.check_and_rename_images_in_folder_with_check(folder_path, naming_patterns)

        return total_renamed

#重新檢查命盟規則是否符合naming pattern狀況
    def check_and_rename_images_in_folder_with_check(self, folder_path, naming_patterns):
        """
        Helper function to check and rename images in a folder, handling existing names.
        """
        files = sorted(os.listdir(folder_path))
        images = [f for f in files if f.lower().endswith(('.png', '.jpg', '.jpeg'))]

        if len(images) != len(naming_patterns):
            messagebox.showwarning("Mismatch Warning", f"路徑：'{folder_path}'\n影像數目不對！")
            return 0

        used_names = set()
        existing_names = []  # Track names that already exist in the folder
        renamed_count = 0

        for image_name in images:
            file_name, file_extension = os.path.splitext(image_name)

            # Skip renaming if already matches the naming pattern
            if file_name in naming_patterns:
                used_names.add(file_name)
                continue

            # Find an unused valid name for renaming
            new_base_name = next((name for name in naming_patterns if name not in used_names), None)
            if new_base_name is None:
                continue

            new_name = new_base_name + ".jpg"
            src = os.path.join(folder_path, image_name)
            dst = os.path.join(folder_path, new_name)

            # Check if the new name already exists
            if os.path.exists(dst):
                existing_names.append(new_name)
                continue

            try:
                os.rename(src, dst)
                used_names.add(new_base_name)
                renamed_count += 1
            except Exception as e:
                messagebox.showerror("Error", f"重命名失敗: {e}")

        # Show warning for existing names
        if existing_names:
            response = messagebox.askyesno(
                "重命名警告",
                f"以下影像已存在於資料夾 '{folder_path}':\n" +
                "\n".join(existing_names) +
                "\n\n是否要覆蓋這些影像？"
            )
            if response:
                for image_name in existing_names:
                    src = os.path.join(folder_path, image_name)
                    dst = os.path.join(folder_path, image_name)
                    try:
                        os.rename(src, dst)
                        renamed_count += 1
                    except Exception as e:
                        messagebox.showerror("Error", f"覆蓋失敗: {e}")

        return renamed_count

    
    def check_and_rename_sixpiece_with_subfolders(self):
        """
        Search for folders with names containing '外箱六面' or 'DUT六面', and rename the files
        inside those folders and their subfolders based on the folder's name with an appended index.
        """
        total_renamed = 0
        sixpiece_keywords = ["外箱六面", "DUT六面"]  # Keywords to search in folder names

        # Traverse directories
        outer_layer_folders = [d for d in os.listdir(self.dir_path) if os.path.isdir(os.path.join(self.dir_path, d))]

        for folder in outer_layer_folders:
            folder_path = os.path.join(self.dir_path, folder)

            # Check if the folder name matches the six-piece pattern
            if any(keyword in folder for keyword in sixpiece_keywords):
                # Rename files in the folder
                total_renamed += self.rename_files_in_sixpiece_folders(folder_path, folder)

                # Search and rename files in subfolders
                subfolders = [
                    os.path.join(folder_path, sub)
                    for sub in os.listdir(folder_path)
                    if os.path.isdir(os.path.join(folder_path, sub)) and sub.isdigit()
                ]

                for subfolder_path in subfolders:
                    total_renamed += self.rename_files_in_sixpiece_folders(subfolder_path, folder)

        return total_renamed


    def rename_files_in_sixpiece_folders(self, folder_path, folder_name):
        """
        Rename files in the given folder with names containing "外箱六面" or "DUT六面" to match
        the folder name (excluding any prefix before the first "_") with an appended index.
        """
        total_renamed = 0
        overwrite_all = False  # Flag to track if the user wants to overwrite all duplicates in this folder
        existing_names = []  # List to collect existing names

        if os.path.isdir(folder_path):
            # Extract folder name after the first "_"
            new_folder_name = "_".join(folder_name.split("_")[1:]) if "_" in folder_name else folder_name
            #新增六面如果真測到有子資料夾則椅子資料夾之數量辨認
            # Check if the folder name contains "外箱六面" or "DUT六面"
            if "外箱六面" in folder_name or "DUT六面" in folder_name:
                # Check the number of images in the main folder
                image_files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f)) and f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))]
                if len(image_files) != 6:
                    messagebox.showwarning("影像數量錯誤", f"請檢查 {folder_path} 影像數量是否正確(6張)，將跳過命名該資料夾影像！")
                    return total_renamed  # Skip renaming if the count is not six

                # Check the number of images in each subfolder
                subfolders = [d for d in os.listdir(folder_path) if os.path.isdir(os.path.join(folder_path, d)) and d.isdigit()]
                for subfolder in subfolders:
                    subfolder_path = os.path.join(folder_path, subfolder)
                    subfolder_images = [f for f in os.listdir(subfolder_path) if os.path.isfile(os.path.join(subfolder_path, f)) and f.lower().endswith(('.png', '.jpg', '.jpeg', '.bmp'))]
                    if len(subfolder_images) != 6:
                        messagebox.showwarning("影像數量錯誤", f"請檢查 {subfolder_path} 影像數量是否正確(6張)，將跳過命名該資料夾影像！")
                        return total_renamed  # Skip renaming if the count is not six

            files = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]
            files.sort()  # Ensure a consistent order for renaming

            for idx, file in enumerate(files, start=1):
                old_file_path = os.path.join(folder_path, file)
                file_extension = os.path.splitext(file)[1]  # Get the file extension
                new_file_name = f"{new_folder_name}({idx}){file_extension}"
                new_file_path = os.path.join(folder_path, new_file_name)

                # Collect existing names
                if os.path.exists(new_file_path):
                    existing_names.append(new_file_name)
                    continue

                # Rename the file
                try:
                    os.rename(old_file_path, new_file_path)
                    total_renamed += 1
                except Exception as e:
                    messagebox.showerror("Error", f"Failed to rename file: {e}")

            # Show warning for existing names
            if existing_names:
                response = messagebox.askyesno(
                    "Existing Names Detected",
                    f"The following images already exist in the folder '{folder_path}':\n" +
                    "\n".join(existing_names) +
                    "\n\nDo you want to overwrite these images?"
                )
                if response:
                    for image_name in existing_names:
                        src = os.path.join(folder_path, image_name)
                        dst = os.path.join(folder_path, image_name)
                        try:
                            os.rename(src, dst)
                            total_renamed += 1
                        except Exception as e:
                            messagebox.showerror("Error", f"Failed to rename file: {e}")

        return total_renamed
#需修正命名總量計算，目前只針對六面命名完總量進行計算

    def rename_images(self):
        """
        Rename images in folders matching specific patterns.
        """
        total_renamed_outer = self.check_and_rename_outer_layer()  # Optional: Call other rename logic
        total_renamed_sixpiece = self.check_and_rename_sixpiece_with_subfolders()
        self.refresh_views()
        #重新命名數量計算需修改
        total_renamed = total_renamed_outer + total_renamed_sixpiece
        messagebox.showinfo("Rename Complete", f"Renamed {total_renamed} images.")

    # Function to insert images into a Word document
    def insert_images_into_word(self):
        def browse_word_file():
            doc_path = filedialog.askopenfilename(filetypes=[("Word Files", "*.docx")])
            word_entry.delete(0, tk.END)
            word_entry.insert(0, doc_path)

        def browse_image_folder():
            folder_path = filedialog.askdirectory()
            image_entry.delete(0, tk.END)
            image_entry.insert(0, folder_path)

        def start_processing():
            doc_path = word_entry.get()
            image_folder = image_entry.get()
            if not doc_path or not image_folder:
                messagebox.showerror("Error", "Please select both a Word file and an image folder.")
                return

            focus_titles = ["4. Test Photograph："]  # Add other specific titles if needed
            self.insert_images_to_document(doc_path, image_folder, focus_titles)
            messagebox.showinfo("Success", "Images have been inserted successfully!")

        # Create a new window for the Word processing functionality
        word_window = tk.Toplevel(self.root)
        word_window.title("Word Table Image Inserter")

        # Word file selection
        tk.Label(word_window, text="Select Word Document:").grid(row=0, column=0, sticky="e")
        word_entry = tk.Entry(word_window, width=60)
        word_entry.grid(row=0, column=1, padx=5, pady=5)
        tk.Button(word_window, text="Browse", command=browse_word_file).grid(row=0, column=2, padx=5, pady=5)

        # Image folder selection
        tk.Label(word_window, text="Select Image Folder:").grid(row=1, column=0, sticky="e")
        image_entry = tk.Entry(word_window, width=60)
        image_entry.grid(row=1, column=1, padx=5, pady=5)
        tk.Button(word_window, text="Browse", command=browse_image_folder).grid(row=1, column=2, padx=5, pady=5)

        # Start button
        tk.Button(word_window, text="Start", command=start_processing, bg="#4CAF50", fg="white").grid(row=2, column=1, pady=10)

    def insert_images_to_document(self, doc_path, image_folder, focus_titles):
        """Insert images into table cells based on sanitized subtitles and matching names."""
        doc = Document(doc_path)
        elements = doc.element.body  # Access document elements
        prev_title = None  # To track the subtitle above a table

        for element in elements:
            if element.tag.endswith("p"):  # Paragraph
                paragraph_text = element.text.strip()
                if paragraph_text:  # If the paragraph has text
                    print(f"Detected subtitle: {paragraph_text}")
                    prev_title = paragraph_text  # Store subtitle

            elif element.tag.endswith("tbl"):  # Table
                if prev_title:
                    sanitized_title = self.sanitize_title(prev_title)

                    # Focus only on specific titles like "4. Test Photograph："
                    if not any(focus_title in prev_title for focus_title in focus_titles):
                        print(f"Skipping table under title: {prev_title}")
                        prev_title = None
                        continue

                    # Locate the corresponding subfolder
                    subfolder = os.path.join(image_folder, sanitized_title)
                    if not os.path.exists(subfolder):
                        print(f"Subfolder not found: {subfolder}")
                        prev_title = None
                        continue

                    print(f"Processing table under title: {prev_title}")
                    table_image_paths = self.find_image_paths(subfolder)

                    for row in element.rows:
                        for cell in row.cells:
                            cell_text = cell.text.strip()
                            if cell_text in table_image_paths:
                                try:
                                    # Insert image into the cell
                                    cell.text = ""
                                    paragraph = cell.paragraphs[0]
                                    run = paragraph.add_run()
                                    image_path = table_image_paths[cell_text]

                                    image = Image.open(image_path)
                                    aspect_ratio = image.width / image.height
                                    run.add_picture(image_path, width=Inches(2.5), height=Inches(2.5 / aspect_ratio))
                                    print(f"Inserted {image_path} into the cell.")
                                except Exception as e:
                                    print(f"Error inserting image: {e}")
                prev_title = None  # Reset title after processing the table

        # Save the modified document
        new_doc_path = doc_path.replace('.docx', '_with_images.docx')
        doc.save(new_doc_path)
        print(f"Document saved as: {new_doc_path}")
        self.open_file(new_doc_path)

    def open_file(self, file_path):
        """Open the file at the given path."""
        os.startfile(file_path)

    def find_image_paths(self, image_folder):
        """Map image names (without extensions) to their full paths."""
        image_paths = {}
        for root, dirs, files in os.walk(image_folder):
            for file in files:
                if file.lower().endswith(('.png', '.jpg', '.jpeg', '.gif', '.bmp')):
                    image_name = os.path.splitext(file)[0]
                    image_paths[image_name] = os.path.join(root, file)
        return image_paths

    def sanitize_title(self, title):
        """Sanitize title by removing invalid characters for folder names."""
        return title.replace("：", "").replace(":", "").strip()

# Main execution
if __name__ == "__main__":
    default_md_file = r"C:\Users\AlvinYT_Liang\Desktop\AutomaticTool\重新命名\ME rename tool\ME_Other_Folder_Structure.md"
    root = TkinterDnD.Tk()  # Create the root window with drag-and-drop enabled
    app = FileExplorerApp(root)  # Pass the root to FileExplorerApp
    root.mainloop()  # Start the main event loop